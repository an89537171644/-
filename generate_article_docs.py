from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path("out")
OUT_DIR.mkdir(parents=True, exist_ok=True)

ARTICLE_TITLE = (
    "ПРОГНОЗИРОВАНИЕ ПРОДОЛЬНЫХ УСИЛИЙ И ПРЕДВАРИТЕЛЬНЫЙ ПОДБОР "
    "СЕЧЕНИЙ СТЕРЖНЕЙ СТАЛЬНЫХ ФЕРМ С ПРИМЕНЕНИЕМ ИСКУССТВЕННОЙ "
    "НЕЙРОННОЙ СЕТИ"
)
JOURNAL = "научно-техническом журнале «Региональная архитектура и строительство», г. Пенза"

AUTHORS = [
    {
        "name": "Николюкин Алексей Николаевич",
        "short": "А.Н. Николюкин",
        "position": "Старший преподаватель кафедры «Конструкции зданий и сооружений»",
        "degree": "к.т.н.",
        "work_phone": "+7 (953) 717-16-44",
        "mobile_phone": "+7 (953) 717-16-44",
        "email": "valax1@yandex.ru",
        "contribution": "25%",
    },
    {
        "name": "Монастырёв Павел Владиславович",
        "short": "П.В. Монастырёв",
        "position": "Директор Института архитектуры, строительства и транспорта, профессор",
        "degree": "д.т.н., профессор; член-корреспондент РААСН",
        "work_phone": "+7 (960) 663-00-93",
        "mobile_phone": "+7 (960) 663-00-93",
        "email": "monastyrev68@mail.ru",
        "contribution": "25%",
    },
    {
        "name": "Hai Fang",
        "short": "Hai Fang",
        "position": "Professor, College of Civil Engineering, Nanjing Tech University",
        "degree": "—",
        "work_phone": "—",
        "mobile_phone": "—",
        "email": "fanghainjut@163.com",
        "contribution": "25%",
    },
    {
        "name": "Xinchen Zhang",
        "short": "Xinchen Zhang",
        "position": "Associate Professor, College of Civil Engineering, Nanjing Tech University",
        "degree": "—",
        "work_phone": "—",
        "mobile_phone": "—",
        "email": "xc.zhang@njtech.edu.cn",
        "contribution": "25%",
    },
]


def set_rfonts(rpr, name: str = "Times New Roman") -> None:
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def style_run(run, *, size: float = 10.5, bold: bool = False, italic: bool = False,
              underline: bool = False, color: RGBColor | None = None,
              font_name: str = "Times New Roman"):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color is not None:
        run.font.color.rgb = color
    set_rfonts(run._element.get_or_add_rPr(), font_name)
    return run


def set_paragraph(paragraph, *, align=None, before: float = 0, after: float = 0,
                  line: float = 1.0, first_line: float | None = None,
                  left: float | None = None, right: float | None = None,
                  keep_with_next: bool = False) -> None:
    fmt = paragraph.paragraph_format
    if align is not None:
        paragraph.alignment = align
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = line
    if first_line is not None:
        fmt.first_line_indent = Cm(first_line)
    if left is not None:
        fmt.left_indent = Cm(left)
    if right is not None:
        fmt.right_indent = Cm(right)
    fmt.keep_with_next = keep_with_next


def set_default_document_style(doc: Document, font_size: float) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(font_size)
    set_rfonts(style._element.get_or_add_rPr(), "Times New Roman")
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.line_spacing = 1.0


def set_a4(section, *, top: float = 1.45, bottom: float = 1.45,
           left: float = 1.55, right: float = 1.55) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_margins(cell, top=35, start=55, bottom=35, end=55) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_text(paragraph, text: str, **kwargs):
    return style_run(paragraph.add_run(text), **kwargs)


def add_underlined_title(doc: Document, title: str, *, size: float = 10.5,
                         align=WD_ALIGN_PARAGRAPH.JUSTIFY, before: float = 0,
                         after: float = 0) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, align=align, before=before, after=after, line=1.0)
    add_text(p, title, size=size, underline=True)


def add_bottom_border(paragraph, *, size: str = "6", color: str = "000000") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def make_logo(path: Path) -> None:
    width, height = 700, 300
    image = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    ring = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    rdraw = ImageDraw.Draw(ring)
    box1 = (205, 30, 500, 265)
    box2 = (235, 42, 470, 253)
    rdraw.ellipse(box1, outline=(45, 45, 45, 255), width=7)
    rdraw.ellipse(box2, outline=(95, 95, 95, 255), width=3)
    ring = ring.rotate(-17, resample=Image.Resampling.BICUBIC, expand=False)
    image.alpha_composite(ring)
    draw = ImageDraw.Draw(image)
    font_candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed-Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ]
    font_path = next((p for p in font_candidates if Path(p).exists()), None)
    font = ImageFont.truetype(font_path, 112) if font_path else ImageFont.load_default()
    text = "ТГТУ"
    bbox = draw.textbbox((0, 0), text, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text(((width - tw) / 2, (height - th) / 2 - 12), text, font=font,
              fill=(10, 10, 10, 255))
    image.save(path)


def build_conclusion() -> Path:
    doc = Document()
    set_default_document_style(doc, 10.0)
    section = doc.sections[0]
    set_a4(section, top=1.25, bottom=1.25, left=1.35, right=1.35)

    approval = doc.add_table(rows=1, cols=2)
    approval.alignment = WD_TABLE_ALIGNMENT.CENTER
    approval.autofit = False
    remove_table_borders(approval)
    approval.columns[0].width = Cm(9.8)
    approval.columns[1].width = Cm(8.2)
    right = approval.cell(0, 1)
    set_cell_margins(right, top=0, start=0, bottom=0, end=0)
    lines = [
        ("УТВЕРЖДАЮ", False, 10.5),
        ("Проректор по научной работе", False, 10.5),
        ("____________________ Д.Ю. Муромцев", False, 10.0),
        ("(подпись, инициалы и фамилия)", False, 7.5),
        ("«____» _______________________ 2026 г.", False, 10.0),
    ]
    for idx, (txt, bold, size) in enumerate(lines):
        p = right.paragraphs[0] if idx == 0 else right.add_paragraph()
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, after=0, line=1.0)
        add_text(p, txt, size=size, bold=bold)

    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=0, line=1.0, keep_with_next=True)
    add_text(p, "ЗАКЛЮЧЕНИЕ", size=15.5, bold=True)
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=2, line=1.0, keep_with_next=True)
    add_text(p, "о возможности открытого опубликования", size=10.5, bold=True)

    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=1, line=1.0)
    add_text(p, "статьи «", size=10.0)
    add_text(p, ARTICLE_TITLE, size=10.0, underline=True)
    add_text(p, "» авторов А.Н. Николюкина, П.В. Монастырёва, Hai Fang, Xinchen Zhang", size=10.0, underline=True)
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=1, line=1.0)
    add_text(p, "(название статьи или монографии, подлежащих экспертизе, инициалы и фамилия авторов)", size=7.3)

    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=0, line=1.0)
    add_text(p, "Экспертная комиссия ФГБОУ ВО «ТГТУ» в составе¹ ", size=10.0)
    add_text(p, "Соколов М.В. — д.т.н., профессор кафедры «Компьютерно-интегрированные системы в машиностроении»; ", size=10.0, underline=True)
    add_text(p, "Ерофеев А.В. — к.т.н., доцент кафедры «Конструкции зданий и сооружений»; ", size=10.0, underline=True)
    add_text(p, "Чепрасова Т.И. — ведущий специалист по защите государственной тайны; ", size=10.0, underline=True)
    add_text(p, "Кузнецова М.С. — начальник отдела патентоведения, стандартизации и метрологии", size=10.0, underline=True)
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0)
    add_text(p, "(фамилия и инициалы членов комиссии, наименование должности)", size=7.3)

    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=0, line=1.0)
    add_text(p, "в период с «____» _______________ 2026 г. по «____» ________________ 2026 г. провела экспертизу материалов рукописи статьи", size=10.0)
    add_underlined_title(doc, f"«{ARTICLE_TITLE}»", size=10.0, after=0)
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0)
    add_text(p, "(название статьи или монографии, подлежащих экспертизе)", size=7.3)

    paragraphs = [
        "на предмет отсутствия (наличия) в них сведений, составляющих государственную тайну, и возможности (невозможности) их открытого опубликования.",
        "Руководствуясь Законом РФ «О государственной тайне», Перечнем сведений, отнесенных к государственной тайне, утвержденным Указом Президента РФ от 30 ноября 1995 г. № 1203, а также Перечнем сведений, подлежащих засекречиванию, утверждённым приказом Минобрнауки России от 04 декабря 2023 года № 31с, комиссия установила:",
        "I. Сведения, содержащиеся в рассматриваемых материалах, находятся в компетенции ФГБОУ ВО «ТГТУ».",
    ]
    for txt in paragraphs:
        p = doc.add_paragraph()
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=0, line=1.0, first_line=0.8)
        add_text(p, txt, size=10.0)

    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=0, line=1.0, first_line=0.8)
    add_text(p, "А) Сведения, содержащиеся в рассматриваемых материалах рукописи статьи ", size=10.0)
    add_text(p, f"«{ARTICLE_TITLE}»", size=10.0, underline=True)
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0)
    add_text(p, "(название статьи или монографии, содержащиеся в материалах)", size=7.3)

    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=0, line=1.0, first_line=0.8)
    add_text(p, "не подпадают под действие Перечня сведений, составляющих государственную тайну (статья 5 Закона Российской Федерации «О государственной тайне»), не относятся к Перечню сведений, отнесенных к государственной тайне, утвержденному Указом Президента Российской Федерации от 30 ноября 1995 г. № 1203, не подлежат засекречиванию и данные материалы могут быть открыто опубликованы.", size=10.0)

    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=1, after=1, line=1.0)
    add_text(p, "__________________________", size=8.0)
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=1, line=1.0)
    add_text(p, "¹ Или руководитель-эксперт, если экспертиза материалов проводится руководителем структурного подразделения организации, в котором работает автор подготовленных материалов.", size=7.2)

    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=1, after=1, line=1.0, keep_with_next=True)
    add_text(p, "Члены комиссии:", size=10.0, bold=True)

    signatures = [
        ("Председатель экспертной комиссии, эксперт по вопросам конструкторско-технологического обеспечения машиностроительных производств", "М.В. Соколов"),
        ("Эксперт по вопросам строительства", "А.В. Ерофеев"),
        ("Эксперт по обеспечению установленного режима секретности в части предотвращения распространения сведений, составляющих государственную тайну, при открытом опубликовании материалов", "Т.И. Чепрасова"),
        ("Секретарь экспертной комиссии, эксперт по вопросам патентной экспертизы", "М.С. Кузнецова"),
    ]
    for role, name in signatures:
        p = doc.add_paragraph()
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, after=0, line=1.0)
        add_text(p, role, size=8.9, underline=True)
        p = doc.add_paragraph()
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.RIGHT, after=0, line=1.0)
        add_text(p, "__________________________________________  ", size=8.8)
        add_text(p, name, size=9.0)
        p = doc.add_paragraph()
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.RIGHT, after=0, line=1.0)
        add_text(p, "(подпись, инициалы и фамилия)", size=6.8)

    doc.add_page_break()
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=640, after=0, line=1.0)
    add_text(p, "Зарегистрировано «____» _______________ 2026 г.        рег. № ____________________", size=10.0)

    path = OUT_DIR / "zaklyuchenie_open_publication.docx"
    doc.save(path)
    return path


def add_author_table(doc: Document, author: dict) -> None:
    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(0.75), Cm(6.1), Cm(9.55)]
    labels = [
        ("1", "Фамилия Имя Отчество\n(полностью) автора", author["name"]),
        ("2", "Должность автора, подразделение", author["position"]),
        ("3", "Ученая степень, ученое звание автора", author["degree"]),
        ("4", "Рабочий телефон", author["work_phone"]),
        ("5", "Мобильный телефон", author["mobile_phone"]),
        ("6", "E-mail", author["email"]),
        ("7", "Вклад автора в создание статьи,\nв процентах", author["contribution"]),
        ("", "Подпись", ""),
    ]
    for row_idx, row_data in enumerate(labels):
        row = table.rows[row_idx]
        set_row_cant_split(row)
        for col_idx, (cell, width, value) in enumerate(zip(row.cells, widths, row_data)):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=18, start=35, bottom=18, end=35)
            p = cell.paragraphs[0]
            set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT if col_idx else WD_ALIGN_PARAGRAPH.CENTER,
                          after=0, line=1.0)
            if row_idx == 6 and col_idx == 2:
                add_text(p, value, size=8.2, color=RGBColor(255, 0, 0), bold=False)
            elif row_idx == 6 and col_idx == 1:
                # Bold only the final phrase as in the source form.
                add_text(p, "Вклад автора в создание статьи,\n", size=8.2)
                add_text(p, "в процентах", size=8.2, bold=True)
            else:
                add_text(p, value, size=8.2)
    spacer = doc.add_paragraph()
    set_paragraph(spacer, after=2, line=1.0)


def build_notification() -> Path:
    logo_path = OUT_DIR / "tstu_logo.png"
    make_logo(logo_path)

    doc = Document()
    set_default_document_style(doc, 10.3)
    section = doc.sections[0]
    set_a4(section, top=1.15, bottom=1.15, left=1.45, right=1.45)

    doc.add_picture(str(logo_path), width=Cm(3.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(doc.paragraphs[-1], after=1, line=1.0)

    header_lines = [
        "Министерство науки и высшего образования Российской Федерации",
        "Федеральное государственное бюджетное образовательное учреждение",
        "высшего образования",
        "«ТАМБОВСКИЙ ГОСУДАРСТВЕННЫЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ»",
    ]
    for line in header_lines:
        p = doc.add_paragraph()
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0, keep_with_next=True)
        add_text(p, line, size=11.0, bold=True)

    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=5, after=0, line=1.0, keep_with_next=True)
    add_text(p, "УВЕДОМЛЕНИЕ", size=15.0, bold=True)
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=5, line=1.0, keep_with_next=True)
    add_text(p, "О СОЗДАНИИ ПРОИЗВЕДЕНИЯ", size=15.0, bold=True)

    for idx, author in enumerate(AUTHORS):
        p = doc.add_paragraph()
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, after=0, line=1.0, left=0.0 if idx == 0 else 4.0)
        if idx == 0:
            add_text(p, "Авторы ", size=10.3)
        add_text(p, author["name"], size=10.3, underline=True)
        p = doc.add_paragraph()
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0)
        add_text(p, "(фамилия, инициалы)", size=7.3)

    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=2, after=0, line=1.0, first_line=0.8)
    add_text(p, "А.Н. Николюкин и П.В. Монастырёв, являясь сотрудниками ФГБОУ ВО «ТГТУ», совместно с иностранными соавторами Hai Fang и Xinchen Zhang представляют на рассмотрение экспертной комиссии по определению возможности открытой публикации статьи", size=10.3)
    add_underlined_title(doc, f"«{ARTICLE_TITLE}»", size=10.3, after=0)
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0)
    add_text(p, "(название статьи или монографии, подлежащих экспертизе)", size=7.3)

    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3, line=1.0, first_line=0.8)
    add_text(p, f"предназначенной для опубликования в {JOURNAL}.", size=10.3)

    declarations = [
        "1. Авторы подтверждают, что рукопись статьи не содержит неопубликованные материалы другой организации без разрешения руководителя этой организации на предмет открытого опубликования.",
        "2. Авторы подтверждают, что рукопись статьи не содержит материалов научно-исследовательских, опытно-конструкторских и иных работ, выполняемых на основе государственных контрактов (договоров, соглашений, заданий, грантов, программ или проектов, по которым осуществляется бюджетное финансирование).",
    ]
    for txt in declarations:
        p = doc.add_paragraph()
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=4, line=1.0, first_line=0.8)
        add_text(p, txt, size=10.3)

    doc.add_page_break()
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, after=4, line=1.0, keep_with_next=True)
    add_text(p, "3. Сведения об авторах (на каждого автора):", size=10.0, bold=True)

    for author in AUTHORS:
        add_author_table(doc, author)

    path = OUT_DIR / "uvedomlenie_creation_article.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    conclusion = build_conclusion()
    notification = build_notification()
    print(f"Generated: {conclusion}")
    print(f"Generated: {notification}")
